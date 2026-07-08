# /// script
# dependencies = [
#   "python-docx",
# ]
# ///

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_shading(cell, color_hex):
    # Set background color of a cell
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # Set padding in twentieths of a point (dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        create_attribute(node, 'w:w', str(val))
        create_attribute(node, 'w:type', 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_document():
    doc = Document()
    
    # Page setup - Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Font Settings
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Microsoft JhengHei' # 微軟正黑體
    style_normal.font.size = Pt(11)
    # Ensure Chinese fonts apply correctly
    style_normal._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), 'Microsoft JhengHei')

    # Title Style
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("新北市永和區秀朗國民小學\n學校特色簡報內容")
    title_run.font.name = 'Microsoft JhengHei'
    title_run._element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), 'Microsoft JhengHei')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138) # Academic Blue (#1E3A8A)
    title.paragraph_format.space_after = Pt(20)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("秀朗特色簡報・共四頁主題架構")
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate Gray (#64748B)
    sub.paragraph_format.space_after = Pt(40)

    # Helper function to add slides
    def add_slide_section(slide_num, title_text, tagline, points, note_text=None):
        # Header / Slide Title
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        run_num = h.add_run(f"第 {slide_num} 頁：")
        run_num.font.size = Pt(16)
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(6, 182, 212) # Aqua Cyan (#06B6D4)
        
        run_title = h.add_run(title_text)
        run_title.font.size = Pt(16)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue (#1E3A8A)
        
        # Tagline / Sub-header
        tag = doc.add_paragraph()
        tag_run = tag.add_run(tagline)
        tag_run.font.size = Pt(12)
        tag_run.font.italic = True
        tag_run.font.color.rgb = RGBColor(100, 116, 139)
        tag.paragraph_format.space_after = Pt(12)

        # Bullet Points
        for point_title, point_desc in points:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            
            bold_run = p.add_run(f"【{point_title}】")
            bold_run.bold = True
            bold_run.font.color.rgb = RGBColor(30, 58, 138)
            
            desc_run = p.add_run(point_desc)
            desc_run.font.size = Pt(11)

        # Optional Highlight Callout Box
        if note_text:
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell = table.cell(0, 0)
            set_cell_shading(cell, "F1F5F9") # Soft slate gray background
            set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
            
            # Left border style (academic blue thick accent border)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>\n'
                f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="1E3A8A"/>\n'
                f'  <w:top w:val="none"/>\n'
                f'  <w:right w:val="none"/>\n'
                f'  <w:bottom w:val="none"/>\n'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            cp_run_bold = cp.add_run("💡 重點提示：")
            cp_run_bold.bold = True
            cp_run_bold.font.color.rgb = RGBColor(245, 158, 11) # Gold Accent (#F59E0B)
            
            cp_run_text = cp.add_run(note_text)
            cp_run_text.font.size = Pt(10.5)
            cp_run_text.font.italic = True
            
            # Add page spacing after the callout
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(12)
            spacer.paragraph_format.space_after = Pt(12)
            
        doc.add_page_break()

    # SLIDE 1
    add_slide_section(
        slide_num=1,
        title_text="學校概況與歷史榮耀 (Guinness World Record)",
        tagline="深厚的歷史底蘊與曾經創下的「世界之最」奇蹟",
        points=[
            ("基本資訊", "秀朗國小創立於民國 65 年（1976年）8 月 2 日，位於新北市永和區得和路 202 號。校名源自當地原住民歷史地名「秀朗社」，意在文化傳承。"),
            ("金氏世界紀錄", "全盛時期（約民國 74-75 學年度），學生人數突破 12,000 人（共 222 班），榮登金氏世界紀錄「全世界學生人數最多的小學」，成為當時聞名國際的「超級小學」，日本 TBS 電視台曾專程來台進行報導。"),
            ("歷史定位", "作為台灣教育史上的指標性大校，它是無數永和在地人的共同成長記憶，更是台灣都市化進程中的歷史縮影。")
        ],
        note_text="全盛時期由於學生人數極為龐大，放學時甚至需要實施分流與分時段離校，是一代人的獨特回憶。"
    )

    # SLIDE 2
    add_slide_section(
        slide_num=2,
        title_text="藝文卓越與體育基石 (Arts & Sports Specialties)",
        tagline="推動適性多元發展，讓每個孩子適得其所",
        points=[
            ("音樂社團實力深厚", "學校長年經營高水準音樂社團，設有完整的管樂團、弦樂團及合唱團，是各類音樂競賽的全國常勝軍。"),
            ("傳統藝文代表：歌仔戲團", "擁有獨特特色的「學生歌仔戲團」，成功將台灣本土歌仔戲文化帶入校園與現代舞台，曾受邀遠赴國外進行交流展演，大獲好評。"),
            ("游泳教育特色搖籃", "配備專業的室內溫水游泳池，並設有體育班深耕水上體育，是新北市及全台重要的基層游泳運動選手培育基地。")
        ],
        note_text="傳統民俗歌仔戲與現代體育游泳訓練的結合，展現出秀朗國小動靜皆宜、文武雙全的卓越教育風格。"
    )

    # SLIDE 3
    add_slide_section(
        slide_num=3,
        title_text="智慧科技創客與閱讀素養 (Smart Tech & Reading)",
        tagline="實踐跨領域跨思維，建構未來世代的素養能力",
        points=[
            ("創客與無人機教育", "學校積極建置創客（Maker）基地，融入前沿無人機科技教育，推廣手作與程式邏輯，引導學童從小建立工程思維與動手解決問題的能力。"),
            ("「明日閱讀」深耕", "長期配合推動晨間自主閱讀，打造安靜專注的閱讀環境，透過大量閱讀培養自主思考與閱讀理解力。"),
            ("科學教育多元遊戲", "引進科學遊戲與科學探究競賽，以寓教於樂的方式啟發孩童對於物理、化學等自然科學的探索興趣。")
        ],
        note_text="科技與人文並進！在動手做（科技創客）與安靜讀（明日閱讀）的環境下，孩子們的左右腦能均衡發展。"
    )

    # SLIDE 4
    add_slide_section(
        slide_num=4,
        title_text="規模轉型與精緻化展望 (Transformation & Outlook)",
        tagline="迎向少子化挑戰，將「人口優勢」轉化為「精緻品質」",
        points=[
            ("學生規模精緻化", "因少子化社會變遷，學生人數由全盛期萬人規模，穩健轉型至目前約 3,500 人左右。班級人數減少，使班級經營更貼心精緻。"),
            ("校園空間活用", "原本應對萬人規模的廣大校園空間被活化再利用，轉化為更完善的專科教室、創客空間、智慧多媒體教室及學生活動展演區。"),
            ("精緻多元願景", "學校以「高品質、精緻化、智慧科技、適性揚才」為發展核心，致力於降低每班學生人數，提升個人化關懷，打造新世代精緻教育的典範校園。")
        ],
        note_text="人數的減少代表每一位學生獲得了更充沛的教育關懷與硬體空間。少子化在秀朗國小成了教育轉型精緻化的絕佳養分。"
    )

    # Remove the final slide's trailing page break
    if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text == "":
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)
    
    # Save the file
    output_filename = "秀朗國小特色簡報內容.docx"
    doc.save(output_filename)
    print(f"Word document successfully created and saved as '{output_filename}'.")

if __name__ == '__main__':
    build_document()
